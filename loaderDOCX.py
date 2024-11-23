from src.ingestion.loaders.loaderBase import LoaderBase
from docx import Document

class LoaderDOCX(LoaderBase):

    def __init__(self, filepath:str):
        self.filepath=filepath

    def extract_metadata(self):
        docx_path = self.filepath
        doc = Document(docx_path)
        
        properties = doc.core_properties
        metadata = {
            "title": properties.title,
            "author": properties.author,
            "subject": properties.subject,
            "keywords": properties.keywords,
            "last_modified_by": properties.last_modified_by,
            "created": properties.created,
            "modified": properties.modified,
            "content_status": properties.content_status,
        }
        
        return metadata
        
    def extract_text(self):
        docx_path = self.filepath
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)